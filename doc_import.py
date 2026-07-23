"""
Tim 学习助手 - 文档导入模块（PDF / DOCX）

扫描版 PDF 会渲染为图片并通过百度 OCR 识别；
文本型 PDF 与 DOCX 直接提取文本，构造与 OCR 同构的行列表后复用 ocr_parser 拆题。
"""

import os
import sys
from typing import List, Dict, Any, Tuple

# 确保项目根目录在 path 中
ROOT = os.path.dirname(os.path.abspath(__file__))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

from ocr_parser import create_ocr_parser

# ---------- 内部行结构（与 OCR 结果同构） ----------

def _make_line(text: str, center_y: int = 0, confidence: float = 1.0) -> Dict[str, Any]:
    """构造与 OCRProvider.recognize() 返回一致的文本行。"""
    return {
        'text': (text or '').strip(),
        'bbox': [[0.0, 0.0], [1.0, 0.0], [1.0, 1.0], [0.0, 1.0]],
        'confidence': float(confidence),
        'center_y': center_y,
        'center_x': 0,
    }


# ---------- DOCX ----------

def docx_to_lines(path: str) -> List[Dict[str, Any]]:
    """从 .docx 文件提取文本行（段落 + 表格）。"""
    import docx  # python-docx

    doc = docx.Document(path)
    lines = []
    cy = 0

    def add(text):
        nonlocal cy
        t = (text or '').strip()
        if t:
            lines.append(_make_line(t, center_y=cy))
            cy += 1000

    # 段落
    for para in doc.paragraphs:
        add(para.text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    row_texts.append(ct)
            if row_texts:
                add(' | '.join(row_texts))

    return lines


# ---------- PDF ----------

TEXT_PDF_MIN_CHARS = 50  # 全文字符低于此阈值视为扫描件


def pdf_to_lines_or_images(path: str) -> Tuple[str, List[Dict[str, Any]], List[str]]:
    """
    解析 PDF 文件。

    返回 (mode, lines, image_paths)：
      mode='text'  → 文本型 PDF，lines 非空，image_paths=[]
      mode='ocr'   → 扫描版 PDF，lines=[], image_paths 为逐页 PNG 路径列表
    """
    import pdfplumber
    import fitz  # PyMuPDF

    lines = []
    raw_text = ''

    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pt = page.extract_text()
                if pt:
                    raw_text += pt + '\n'
    except Exception:
        raw_text = ''

    # 判断文本量
    chars = len(raw_text.strip())
    if chars >= TEXT_PDF_MIN_CHARS:
        # 文本型：按行构造
        cy = 0
        for line in raw_text.split('\n'):
            t = line.strip()
            if t:
                lines.append(_make_line(t, center_y=cy))
                cy += 1000
        return ('text', lines, [])

    # 扫描件：逐页渲染为 PNG
    from config import Config
    os.makedirs(Config.OCR_TEMP_FOLDER, exist_ok=True)
    image_paths = []

    try:
        fm = fitz.open(path)
        for i, page in enumerate(fm):
            pix = page.get_pixmap(dpi=150)
            out = os.path.join(Config.OCR_TEMP_FOLDER, f'doc_page_{os.getpid()}_{i}.png')
            pix.save(out)
            image_paths.append(out)
        fm.close()
    except Exception:
        pass

    return ('ocr', [], image_paths)


# ---------- 统一解析 ----------

def lines_to_questions(lines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    将结构化行列表经 ocr_parser 拆题，转为前端 save-batch 兼容的字典列表。

    每条字典含：xueke, timu, xueshengdaan, zhengquedaan, cuowufenxi,
    zhishidian, difficulty, confidence, field_confidences。
    """
    if not lines:
        return []

    parser = create_ocr_parser()
    parsed = parser.parse(lines)

    result = []
    for pq in parsed:
        result.append({
            'xueke': pq.xueke,
            'timu': pq.timu,
            'xueshengdaan': pq.xueshengdaan,
            'zhengquedaan': pq.zhengquedaan,
            'cuowufenxi': pq.cuowufenxi,
            'zhishidian': pq.zhishidian,
            'difficulty': 3,
            'confidence': pq.confidence,
            'field_confidences': pq.field_confidences,
        })

    return result


# ---------- 扩展名校验 ----------

ALLOWED_DOC_EXTENSIONS = {'pdf', 'docx', 'doc'}

DOCX_ONLY_EXTENSIONS = {'docx', 'pdf'}


def check_doc_ext(filename: str) -> str:
    """校验文档扩展名，返回小写扩展名；若不支持则返回空字符串。"""
    if '.' not in filename:
        return ''
    ext = filename.rsplit('.', 1)[1].lower()
    if ext not in ALLOWED_DOC_EXTENSIONS:
        return ''
    return ext
